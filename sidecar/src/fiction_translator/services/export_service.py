"""Export service for chapters and projects."""
from __future__ import annotations

import os
from pathlib import Path
from datetime import datetime
from sqlalchemy.orm import Session
from fiction_translator.db.models import Chapter, Segment, Translation, Export, Project


def export_chapter_txt(db: Session, chapter_id: int, target_language: str = "en") -> dict:
    """Export a chapter as plain text."""
    chapter = db.query(Chapter).filter(Chapter.id == chapter_id).first()
    if not chapter:
        raise ValueError(f"Chapter {chapter_id} not found")

    segments = db.query(Segment).filter(
        Segment.chapter_id == chapter_id
    ).order_by(Segment.order).all()

    lines = [f"# {chapter.title}\n"]
    for seg in segments:
        translation = db.query(Translation).filter(
            Translation.segment_id == seg.id,
            Translation.target_language == target_language,
        ).first()
        text = translation.translated_text if translation and translation.translated_text else seg.source_text
        lines.append(text)

    content = "\n".join(lines)

    # Save to file
    export_dir = Path.home() / ".fiction-translator" / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{chapter.title}_{target_language}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    filepath = export_dir / filename
    filepath.write_text(content, encoding="utf-8")

    # Record export
    export = Export(
        chapter_id=chapter_id,
        project_id=chapter.project_id,
        format="txt",
        file_path=str(filepath),
    )
    db.add(export)
    db.commit()

    return {"path": str(filepath), "format": "txt", "size": len(content)}


def export_chapter_docx(db: Session, chapter_id: int, target_language: str = "en") -> dict:
    """Export a chapter as DOCX."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed. Install with: uv add python-docx")

    chapter = db.query(Chapter).filter(Chapter.id == chapter_id).first()
    if not chapter:
        raise ValueError(f"Chapter {chapter_id} not found")

    doc = Document()
    doc.add_heading(chapter.title, 0)

    segments = db.query(Segment).filter(
        Segment.chapter_id == chapter_id
    ).order_by(Segment.order).all()

    for seg in segments:
        translation = db.query(Translation).filter(
            Translation.segment_id == seg.id,
            Translation.target_language == target_language,
        ).first()
        text = translation.translated_text if translation and translation.translated_text else seg.source_text

        if seg.segment_type == "dialogue" and seg.speaker:
            doc.add_paragraph(f"{seg.speaker}: {text}")
        else:
            doc.add_paragraph(text)

    export_dir = Path.home() / ".fiction-translator" / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{chapter.title}_{target_language}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = export_dir / filename
    doc.save(str(filepath))

    # Record export
    export = Export(
        chapter_id=chapter_id,
        project_id=chapter.project_id,
        format="docx",
        file_path=str(filepath),
    )
    db.add(export)
    db.commit()

    return {"path": str(filepath), "format": "docx"}


def list_exports(db: Session, project_id: int) -> list[dict]:
    """List all exports for a project."""
    exports = db.query(Export).filter(
        Export.project_id == project_id
    ).order_by(Export.created_at.desc()).all()
    return [{
        "id": e.id,
        "chapter_id": e.chapter_id,
        "project_id": e.project_id,
        "format": e.format,
        "file_path": e.file_path,
        "created_at": e.created_at.isoformat() if e.created_at else None,
    } for e in exports]


def delete_export(db: Session, export_id: int) -> dict:
    """Delete an export record and optionally the file."""
    export = db.query(Export).filter(Export.id == export_id).first()
    if not export:
        raise ValueError(f"Export {export_id} not found")

    # Optionally delete the physical file
    file_path = Path(export.file_path)
    file_deleted = False
    if file_path.exists():
        file_path.unlink()
        file_deleted = True

    db.delete(export)
    db.commit()
    return {"deleted": True, "id": export_id, "file_deleted": file_deleted}
